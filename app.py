from fastapi import FastAPI
from langchain.prompts import ChatPromptTemplate
from langchain_openai import ChatOpenAI
#from langserve import add_routes
from langchain.document_loaders import DirectoryLoader
from langchain_community.document_loaders import DirectoryLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_core.messages.base import BaseMessage
from fastapi import UploadFile, File, HTTPException
from langchain_openai import OpenAIEmbeddings
from langchain.chains import LLMChain
from enum import Enum
from typing import Any
import uuid
import pylibmagic
from fastapi import Query
import random
import io
import docx
import shutil
import uvicorn
from typing import List, Dict
import numpy as np
import os
from langchain_community.document_loaders import WebBaseLoader
from langchain.chains.combine_documents.stuff import StuffDocumentsChain
from langchain_pinecone import PineconeVectorStore
from langchain_community.retrievers import PineconeHybridSearchRetriever
from fastapi.responses import PlainTextResponse
from pinecone import Pinecone, ServerlessSpec
from langchain_core.runnables import RunnableParallel
from langchain.chains import RetrievalQA
from langchain_openai import OpenAI
from langchain_core.prompts import PromptTemplate
from langchain_core.runnables import RunnablePassthrough
from langchain_core.output_parsers import StrOutputParser
from langchain_core.runnables import Runnable
from pydantic import BaseModel, Field, validator, ValidationError
from langchain.chains import create_history_aware_retriever, create_retrieval_chain
from langchain.chains.combine_documents import create_stuff_documents_chain
from langchain_core.prompts import MessagesPlaceholder
from langchain_core.messages import HumanMessage, AIMessage, SystemMessage
from tqdm.auto import tqdm
from langchain_pinecone import PineconeVectorStore
from langchain_core.documents import Document
from dotenv import load_dotenv
from fastapi.responses import RedirectResponse
import glob
from fastapi import  UploadFile
from fastapi.responses import JSONResponse
from fastapi.responses import StreamingResponse
from fastapi import Depends
from uuid import uuid4
import time
from fastapi.middleware.cors import CORSMiddleware
from fastapi import  Request
from fastapi import  BackgroundTasks, UploadFile, Form
import getpass
import os
import concurrent.futures
from fastapi import  UploadFile
from PyPDF2 import PdfReader  
import docx
import json
from crawl4ai import AsyncWebCrawler
from rank_bm25 import BM25Okapi
from nltk.tokenize import word_tokenize
import nltk

nltk.download('punkt')
nltk.download('wordnet')
nltk.download('omw-1.4')
nltk.download('punkt_tab')
nltk.download('averaged_perceptron_tagger_eng')

class MyModel(BaseModel):
    message: BaseMessage
class Config:
    arbitrary_types_allowed = True

app = FastAPI(
    title="LangChain Server",
    version="version:vac0.1",
    description="")

origins = ["*"]
app.add_middleware(
    CORSMiddleware,
    allow_origins=origins,
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"])

@app.middleware("http")
async def add_process_time_header(request: Request, call_next):
    start_time = time.perf_counter()
    response = await call_next(request)
    process_time = time.perf_counter() - start_time
    response.headers["Processing-Time"] = str(process_time)
    return response

######################### keys #############################
load_dotenv()
openai_api_key = os.getenv("OPENAI_API_KEY")
pinecone_api_key = os.getenv("PINECONE_API_KEY")
pinecone_index_name = os.getenv("PINECONE_INDEX_NAME")

if not openai_api_key:
    raise EnvironmentError("OPENAI_API_KEY is not set. Please set it in your environment.")
if not pinecone_api_key:
    raise EnvironmentError("PINECONE_API_KEY is not set. Please set it in your environment.")
if not pinecone_index_name:
    raise EnvironmentError("PINECONE_INDEX_NAME is not set. Please set it in your environment.")
if openai_api_key:
    os.environ["OPENAI_API_KEY"] = openai_api_key
if pinecone_api_key:
    os.environ["PINECONE_API_KEY"] = pinecone_api_key
if pinecone_index_name:
    os.environ["PINECONE_INDEX_NAME"] = pinecone_index_name
if not os.getenv("PINECONE_API_KEY"):
    os.environ["PINECONE_API_KEY"] = getpass.getpass("Enter your Pinecone API key: ")

############### vectordatabase initializations ####################
pinecone_api_key = os.environ.get("PINECONE_API_KEY")
pinecone_index_name = os.environ.get("PINECONE_INDEX_NAME")
pc = Pinecone(api_key=pinecone_api_key)
pinecone_client = Pinecone(api_key=pinecone_api_key)
def initialize_pinecone_index(client: Pinecone, index_name: str):
    existing_indexes = [index_info["name"] for index_info in client.list_indexes()]
    if index_name not in existing_indexes:
        client.create_index(
            name=index_name,
            dimension=3072,
            metric="cosine",
            spec=ServerlessSpec(cloud="aws", region="us-east-1"))
        while not client.describe_index(index_name).status["ready"]:
            time.sleep(1)
    return client.Index(index_name)
index = initialize_pinecone_index(pinecone_client, pinecone_index_name)

############################# upload #############################
@app.post("/api/artificial-intelligence/upload")
async def upload_files(
    tenantId: str = Form(...),
    files: List[UploadFile] = File(...),
    background_task: BackgroundTasks = BackgroundTasks()):
    """
    Upload and process multiple document files
        PDF, DOCX, and TXT files
    Args:
        tenantId: Identifier for the tenant
        files: List of files to upload
        background_task: Background task handler
    Returns:
        JSON response indicating upload status
    Raises:
        HTTPException: For invalid files or duplicates
    """
    dir_name = str(uuid4())
    os.makedirs(dir_name, exist_ok=True)
    allowed_extensions = {".pdf", ".docx", ".txt"}
    fileName = set()
    for file in files:
        if file.filename in fileName:
            raise HTTPException(
                status_code=400, detail=f"Duplicate file detected: {file.filename}")
        fileName.add(file.filename)
        _, extension = os.path.splitext(file.filename)
        if extension.lower() not in allowed_extensions:
            raise HTTPException(
                status_code=400,
                detail=f"Invalid file type: {file.filename}. Only PDF, DOCX, and TXT files are allowed.",)
        destination = os.path.join(dir_name, file.filename)
        print('creating doc ' + destination)
        with open(destination, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
    docs = load_docs(dir_name, tenantId)
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
    vectorstore = PineconeVectorStore(index_name=pinecone_index_name, embedding=embeddings)
    chunk_size = 800
    chunk_overlap = 150
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    split_docs = []
    chunk_ids = []
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"
    if not os.path.exists(uploaded_documents_path):
        with open(uploaded_documents_path, "w") as f:
            json.dump({}, f)
    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)
    for doc in docs:
        curr_split_docs = text_splitter.split_documents([doc])
        document_id = str(uuid4())
        curr_chunk_ids = [f"{document_id}_chunk_{i+1}" for i in range(len(curr_split_docs))]
        split_docs = split_docs + curr_split_docs
        chunk_ids = chunk_ids + curr_chunk_ids
        upload_documents[document_id] =  {"fileName": doc.metadata['filename'], "id": document_id, "tenantId": tenantId}
    vectorstore.add_documents(documents=split_docs, ids=chunk_ids)
    with open(uploaded_documents_path, "w") as f:
        json.dump(upload_documents, f, indent=4)
    shutil.rmtree(dir_name)
    return {"status": "success", "message": "Files uploaded successfully."}
def load_docs(directory, tenantId):
    loader = DirectoryLoader(directory)
    docs = loader.load()
    for doc in docs:
        doc.metadata['tenantId'] = tenantId
        doc.metadata['filename'] = os.path.basename(doc.metadata['source'])
        print("ooooooooo" + doc.metadata['filename'])
    return docs

############################# links@upload #############################
class UploadLinkRequest(BaseModel):
    tenantId: str
    url: str 
    name: str

class DocumentType(str, Enum):
    FILE = "File"
    LINK = "Link"

@app.post("/api/artificial-intelligence/links")
async def upload_web_url(request: UploadLinkRequest, background_tasks: BackgroundTasks):
    """
    Process and store web URL content.
    
    Args:
        request: UploadLinkRequest containing tenantId and url.
    
    Returns:
        JSON response indicating processing status.
    
    Raises:
        HTTPException: When no valid url are provided.
    """
    all_docs = []  
    urls = request.url.split(",")  

    for url in urls:
        url = url.strip()
        try:
            async with AsyncWebCrawler(verbose=True) as crawler:
                result = await crawler.arun(url=url)  # Running the crawler on the URL
                extracted_content = result.markdown  # Extracted content in markdown format
                doc = Document(page_content=extracted_content, metadata={"source": url, "fileName":request.name, "tenantId": request.tenantId,"type": DocumentType.LINK })
                all_docs.append(doc)

        except Exception as e:
            print(f"Error loading URL {url}: {e}")
            continue

    if not all_docs:
        raise HTTPException(status_code=400, detail="No valid URL provided or failed to fetch content.")
    
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
    vectorstore = PineconeVectorStore(index_name=os.getenv("PINECONE_INDEX_NAME"), embedding=embeddings)
    chunk_size = 13500
    chunk_overlap = 1000
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    split_docs = []
    chunk_ids = []
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"

    if not os.path.exists(uploaded_documents_path):
        with open(uploaded_documents_path, "w") as f:
            json.dump({}, f)
    
    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)
    
    document_id = str(uuid.uuid4())
    for doc in all_docs:
        curr_split_docs = text_splitter.split_documents([doc])
        curr_chunk_ids = [f"{document_id}_chunk_{i+1}" for i in range(len(curr_split_docs))]
        split_docs.extend(curr_split_docs)
        chunk_ids.extend(curr_chunk_ids)
        upload_documents[document_id] = {"source": doc.metadata["source"],"fileName":request.name, "id": document_id, "tenantId": request.tenantId, "type": DocumentType.LINK}
    
    vectorstore.add_documents(documents=split_docs, ids=chunk_ids)
    
    with open(uploaded_documents_path, "w") as f:
        json.dump(upload_documents, f, indent=4)
    
    return {"status": "success", "message": "URL processed and stored successfully.", "data": {"documentId": document_id}}
############################# links@edit #############################
class UpdateLinkRequest(BaseModel):
    tenantId: str
    embeddedId: str
    url: str 
    name: str

# class DocumentType(str, Enum):
#     FILE = "File"
#     LINK = "Link"    
@app.put("/api/artificial-intelligence/links")
async def update_web_url(request: UpdateLinkRequest):
    """
    Update web URL content for a specific embeddedId.
    Args:
        request: UpdateLinkRequest containing tenantId, embeddedId, and url.
    Returns:
        JSON response indicating processing status.
    Raises:
        HTTPException: When no valid url are provided or if the embeddedId does not exist.
    """
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"
    if not os.path.exists(uploaded_documents_path):
        with open(uploaded_documents_path, "w") as f:
            json.dump({}, f)
    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)
    if request.embeddedId not in upload_documents:
        raise HTTPException(status_code=404, detail="The embeddedId does not exist.")
    try:
        all_docs = []  
        urls = request.url.split(",")  

        for url in urls:
            url = url.strip()
            try:
                async with AsyncWebCrawler(verbose=True) as crawler:
                    result = await crawler.arun(url=url)  # Running the crawler on the URL
                    extracted_content = result.markdown  # Extracted content in markdown format
                    doc = Document(page_content=extracted_content, metadata={"source": url,"fileName":request.name, "tenantId": request.tenantId, "type": DocumentType.LINK})
                    all_docs.append(doc)

            except Exception as e:
                print(f"Error loading URL {url}: {e}")
                continue

        if not all_docs:
            raise HTTPException(status_code=400, detail="No valid URL provided or failed to fetch content.")
        
        embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
        vectorstore = PineconeVectorStore(index_name=pinecone_index_name, embedding=embeddings)
        chunk_size = 13500
        chunk_overlap = 1000
        text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap, add_start_index=True)
        split_docs = []
        chunk_ids = []
        
        existing_chunk_ids = [key for key in upload_documents if upload_documents[key]['id'] == request.embeddedId]
        
        for chunk_id in existing_chunk_ids:
            del upload_documents[chunk_id]  # Removing old chunks
            
        ids_to_delete = [id for id in index.list(prefix=request.embeddedId)]
        index.delete(ids=ids_to_delete)

        for doc in all_docs:
            curr_split_docs = text_splitter.split_documents([doc])
            document_id = request.embeddedId  # same embeddedId
            curr_chunk_ids = [f"{document_id}_chunk_{i+1}" for i in range(len(curr_split_docs))]
            split_docs.extend(curr_split_docs)
            chunk_ids.extend(curr_chunk_ids)
            upload_documents[document_id] = {"source": doc.metadata["source"],"id": document_id, "fileName":request.name, "tenantId": request.tenantId, "type": DocumentType.LINK}

        vectorstore.add_documents(documents=split_docs, ids=chunk_ids)
        with open(uploaded_documents_path, "w") as f:
            json.dump(upload_documents, f, indent=4)
        return {
            "status": "success",
            "message": "URL updated successfully.",
            "embeddedId": request.embeddedId,
            "type": DocumentType.LINK,
            "updated_url": request.url
        }   
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error updating url: {str(e)}")
############################# links@delete #############################
class DeleteRequest(BaseModel):
    embeddedId: str
@app.delete("/api/artificial-intelligence/links", summary="Delete Documents", description="Delete all url documents in the Pinecone index that match the given prefix.")
async def delete_url(request: DeleteRequest):
    """Delete url documents matching prefix
    Args:
        request: Delete request containing prefix
    Returns:
        JSON response indicating deletion status
    Raises:
        HTTPException: When no matching url documents found
    """
    prefix = request.embeddedId
    ids_to_delete = [id for id in index.list(prefix=prefix)]
    print(ids_to_delete)
    if not ids_to_delete:
        raise HTTPException(status_code=404, detail="No url documents found with the given prefix.")
    index.delete(ids=ids_to_delete)
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"
    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)
    if prefix in upload_documents:
        del upload_documents[prefix]
    with open(uploaded_documents_path, "w") as f:
        json.dump(upload_documents, f, indent=4)
    return {"message": f"Deleted {len(ids_to_delete)} documents with embeddedId '{prefix}'."}      
############################# qa@upload #############################
class QARequest(BaseModel):
    tenantId: str
    question: str
    answer: str
@app.post("/api/artificial-intelligence/qa")
async def upload_question_answer(
    request: QARequest,
    background_task: BackgroundTasks):
    """
    Upload and store a question-answer pair
    Args:
        request: QA request containing question and answer
        background_task: Background task handler    
    Returns:
        JSON response with QA details and status  
    Raises:
        HTTPException: For processing errors
    """
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
    vectorstore = PineconeVectorStore(index_name=pinecone_index_name, embedding=embeddings)
    llm = ChatOpenAI(model="gpt-3.5-turbo", temperature=0.0)
    chunk_size = 2000
    chunk_overlap = 100
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    split_docs = []
    chunk_ids = []
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"
    if not os.path.exists(uploaded_documents_path):
        with open(uploaded_documents_path, "w") as f:
            json.dump({}, f)
    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)
    try:
        document_id = str(uuid.uuid4())
        combined_text = f"Question: {request.question}\nAnswer: {request.answer}"
        metadata = {
            "tenantId": request.tenantId,
            "fileName": f"Q&A_input_{document_id}.txt",
            "id": document_id,
            "question": request.question,  
            "answer": request.answer}
        document = Document(page_content=combined_text, metadata=metadata)
        curr_split_docs = text_splitter.split_documents([document])
        curr_chunk_ids = [f"{document_id}_chunk_{i+1}" for i in range(len(curr_split_docs))]
        split_docs += curr_split_docs
        chunk_ids += curr_chunk_ids
        upload_documents[document_id] = metadata
        vectorstore.add_documents(documents=split_docs, ids=chunk_ids)
        with open(uploaded_documents_path, "w") as f:
            json.dump(upload_documents, f, indent=4)
        return {
            "status": "success",
            "message": "Question and Answer processed successfully.",
            "id": document_id,
            "question": request.question,  
            "answer": request.answer }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing the question and answer: {str(e)}")      
############################# qa@edit #############################
class QAEditRequest(BaseModel):
    tenantId: str
    embeddedId: str
    question: str
    answer: str
@app.put("/api/artificial-intelligence/qa")
async def edit_question_answer(
    request: QAEditRequest,  
    background_task: BackgroundTasks = BackgroundTasks()):
    """
    Edit a question and answer pair by replacing the content in the same embeddedId.
    """
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
    vectorstore = PineconeVectorStore(index_name=pinecone_index_name, embedding=embeddings)
    chunk_size = 2000
    chunk_overlap = 100
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"
    if not os.path.exists(uploaded_documents_path):
        with open(uploaded_documents_path, "w") as f:
            json.dump({}, f)
    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)
    if request.embeddedId not in upload_documents:
        raise HTTPException(status_code=404, detail="The embeddedId does not exist.")
    try:
        ids_to_delete = [id for id in index.list(prefix=request.embeddedId)]
        index.delete(ids=ids_to_delete)
        combined_text = f"Question: {request.question}\nAnswer: {request.answer}"
        metadata = {
            "tenantId": request.tenantId,
            "fileName": f"Q&A_input_{request.embeddedId}.txt",
            "id": request.embeddedId,
            "question": request.question,
            "answer": request.answer}
        document = Document(page_content=combined_text, metadata=metadata)
        curr_split_docs = text_splitter.split_documents([document])
        curr_chunk_ids = [f"{request.embeddedId}_chunk_{i+1}" for i in range(len(curr_split_docs))]
        vectorstore.add_documents(documents=curr_split_docs, ids=curr_chunk_ids)
        upload_documents[request.embeddedId] = metadata
        with open(uploaded_documents_path, "w") as f:
            json.dump(upload_documents, f, indent=4)
        return {
            "status": "success",
            "message": "Question and Answer edited successfully.",
            "id": request.embeddedId,
            "question": request.question,
            "answer": request.answer }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error editing the question and answer: {str(e)}")

####################################### qa@delete ##########################
class DeleteRequest(BaseModel):
    prefix: str
@app.delete("/api/artificial-intelligence/qa", summary="Delete Question Answer", description="Delete all documents in the Pinecone index that match the given prefix.")
async def delete_documents(request: DeleteRequest):
    """Delete Question Answer matching prefix
    Args:
        request: Delete request containing prefix
    Returns:
        JSON response indicating deletion status
    Raises:
        HTTPException: When no matching documents found
    """
    prefix = request.prefix
    ids_to_delete = [id for id in index.list(prefix=prefix)]
    print(ids_to_delete)
    if not ids_to_delete:
        raise HTTPException(status_code=404, detail="No QA found with the given prefix.")
    index.delete(ids=ids_to_delete)
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"
    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)
    if prefix in upload_documents:
        del upload_documents[prefix]
    with open(uploaded_documents_path, "w") as f:
        json.dump(upload_documents, f, indent=4)
    return {"message": f"Deleted {len(ids_to_delete)} QA with prefix '{prefix}'."}


############################# tenant_files #############################
# @app.get("/api/artificial-intelligence/tenant_files")
# async def retrieve_files(tenantId: str = Query(...)):
#     uploaded_documents_path = "do_not_delete_uploaded_documents.json"
#     if not os.path.exists(uploaded_documents_path):
#         with open(uploaded_documents_path, "w") as f:
#             json.dump({}, f)
#     with open(uploaded_documents_path, "r") as f:
#         upload_documents = json.load(f)
#     tenantFiles = []
#     for item in list(upload_documents.values()):
#         if item.get('tenantId') == tenantId:
#             tenantFiles.append(item)
#     return {"data": tenantFiles}

############################# tenant_files #############################
@app.get("/api/artificial-intelligence/tenant_files")
async def retrieve_files(tenantId: str = Query(...)):
    print("Fetching uploaded file lists")
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"

    if not os.path.exists(uploaded_documents_path):
        with open(uploaded_documents_path, "w") as f:
            json.dump({}, f)

    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)

    tenantFiles = []
    for item in upload_documents.values():
        if item.get('tenantId') == tenantId:
            if 'fileName' in item and not item.get('fileName', '').startswith('Q&A_input_'):
                tenantFiles.append(item)
            elif 'source' in item and item['source'].startswith('https'):
                tenantFiles.append(item)

    return {"data": tenantFiles}

############################# prompts #############################
def initializeVectorStore():
    """Initialize and configure vector store with embeddings
    Returns:
        Configured PineconeVectorStore instance
    """
    embeddings = OpenAIEmbeddings(model="text-embedding-ada-002")
    vectorstore = PineconeVectorStore(index_name=pinecone_index_name, embedding=embeddings)
    return vectorstore

llm = ChatOpenAI(
    model="gpt-4o-mini",
    temperature=0.0,
)

def detect_request_type(keyword, vectorstore, tenantId):
    """Detect the type of request based on the keyword and metadata of the documents."""
    retriever = vectorstore.as_retriever(
        search_type="similarity",
        search_kwargs={
            "k": 5,
            "filter": {
                'tenantId': {'$eq': tenantId}
            }
        }
    )
    docs = retriever.get_relevant_documents(keyword)
    
    #  metadata to determine the type
    tenant_files = []
    for doc in docs:
        metadata = doc.metadata  
        if metadata.get('tenantId') == tenantId:
            if 'fileName' in metadata and not metadata.get('fileName', '').startswith('Q&A_input_'):
                tenant_files.append(metadata)
            elif 'source' in metadata and metadata['source'].startswith('https'):
                tenant_files.append(metadata)
    
    #  type based on metadata
    if any(metadata.get('fileName', '').startswith('Q&A_input_') for metadata in tenant_files):
        return "qa"
    elif any(metadata.get('source', '').startswith('https') for metadata in tenant_files):
        return "Link"
    else:
        return "Document"

@app.get("/api/artificial-intelligence/prompts")
async def unified_prompts_endpoint(
    tenantId: str = Query(...), 
    keyword: str = Query(...)
):
    """Unified endpoint to handle different types of prompts
    Args:
        tenantId: Identifier for the tenant
        keyword: Search keyword/prompt
    Returns:
        PlainText response with answer or error message
    """
    try:
        print(tenantId, keyword)
        vectorstore = initializeVectorStore()
        
        # the type of request
        request_type = detect_request_type(keyword, vectorstore, tenantId)
        
        if request_type == "Document":
            template = """Answer the question based only on the following context:
            {context}
            User Question:
            {question}
            If the user asks you, 'Who are you?' or 'What are you?' respond with: "I am an AI assistant."
            If the user greets you, respond with one of the following without searching information from knowledge base:
            - "Hello! How can I assist you today?"
            - "Hi there! What can I do for you?"
            - "Good day! How can I assist you?"
            - "Hey! What can I help you with?"
            - "Hi! How can I support you today?"
            - "Welcome! How may I help you?"
            - "Salutations! How can I assist you?"
            - "What's up? How can I help?"
            - "Good to see you! How can I assist?"
            Please provide a concise and accurate answer based on the context above. If the context does not contain information to answer the question, respond with: "I don't have the information you're looking for, please provide additional details."
            Make sure your answer is relevant and accurate to the question and does not repeat the question itself.
            In response give only answer but not question.
            """
            retriever = vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={
                    "k": 5,
                    "filter": {
                        'tenantId': {'$eq': tenantId}
                    }
                }
            )
            prompt = ChatPromptTemplate.from_template(template)
            rag_chain = (
                {"context": retriever, "question": RunnablePassthrough()}
                | prompt
                | llm
                | StrOutputParser())
            answer = rag_chain.invoke(keyword)
            return PlainTextResponse(answer)

        elif request_type == "Link":
            template = """Based on the provided web content, answer the following question:
            {context}
            User Question:
            {question}
            If the user asks you, 'Who are you?' or 'What are you?' respond with: "I am an AI assistant."
            If the user greets you, respond with one of the following without searching information from knowledge base:
            - "Hello! How can I assist you today?"
            - "Hi there! What can I do for you?"
            - "Good day! How can I assist you?"
            - "Hey! What can I help you with?"
            - "Hi! How can I support you today?"
            - "Welcome! How may I help you?"
            - "Salutations! How can I assist you?"
            - "What's up? How can I help?"
            - "Good to see you! How can I assist?"
             Instructions for answering:
            1. Extract relevant information directly from the web content
            2. If the question is about specific data or facts, cite them accurately
            3. If the content contains multiple relevant sections, synthesize them into a coherent answer
            4. For technical or complex topics, provide clear, simplified explanations
            5. If the information is not available in the context, respond with: "This information is not available in the source content. Please try a different question or provide more details."
            6. For greetings, use one of the greeting responses above
            7. If the user asks questions that are synonyms or rephrased versions of questions in URL links, provide relevant answers based on the intent of the question, not just the exact wording.
            8. When answering, focus on matching the intent of the user's question, even if the exact words are not present in the context. Use contextual understanding to infer the meaning and provide accurate responses.
            9. If the question is ambiguous or unclear, ask for clarification to ensure the response aligns with the user's intent.
            10. For numerical data or specific metrics (e.g., measurements, statistics, or specifications), provide the exact match from the context.
            Remember:- For example: The secondary sensor of the Google Pixel 9 with 48 MP ultra-wide-angle, f2.2 and field of view (FOV) is 125.8˚and
                                 The secondary sensor of the Galaxy S25 with 12 MP ultra-wide-angle, f/2.2 aperture and field of view (FOV) is 120˚
            11. If the user asks any date related information, search the context for date and provide the  match.
                The date extraction should work with formats like:    
                - 2022/02/09
                - 04/03/2024
                - 01/07/2023
                - 2022, January 09
                - March 09, 2022
                Remember:- For example: The Samsung Galaxy S22 Ultra 5G was announced in 2022, February 09.
            """
            retriever = vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={
                    "k": 5,  
                    "filter": {
                        'tenantId': {'$eq': tenantId}
                    }
                }
            )
            docs = retriever.get_relevant_documents(keyword)
            corpus = [doc.page_content for doc in docs]
            tokenized_corpus = [word_tokenize(doc.lower()) for doc in corpus]
            bm25 = BM25Okapi(tokenized_corpus)
            tokenized_query = word_tokenize(keyword.lower())
            bm25_scores = bm25.get_scores(tokenized_query)
            reranked_pairs = list(zip(docs, bm25_scores))
            reranked_pairs.sort(key=lambda x: x[1], reverse=True)
            reranked_docs = [pair[0] for pair in reranked_pairs[:1]]
            prompt = ChatPromptTemplate.from_template(template)
            rag_chain = (
                {"context": lambda x: reranked_docs, "question": RunnablePassthrough()}
                | prompt
                | llm
                | StrOutputParser())
            answer = rag_chain.invoke(keyword)
            return PlainTextResponse(answer)

        elif request_type == "qa":
            template = """Answer the question based only on the following context:
            {context}
            User Question:
            {question}
            If the user asks you, 'Who are you?' or 'What are you?' respond with: "I am an AI assistant."
            If the user greets you, respond with one of the following without searching information from knowledge base:
            - "Hello! How can I assist you today?"
            - "Hi there! What can I do for you?"
            - "Good day! How can I assist you?"
            - "Hey! What can I help you with?"
            - "Hi! How can I support you today?"
            - "Welcome! How may I help you?"
            - "Salutations! How can I assist you?"
            - "What's up? How can I help?"
            - "Good to see you! How can I assist?"
            Please follow these rules:
            1. If the user's question EXACTLY matches a question in the context (case-insensitive), provide the corresponding answer.
            2. If there is no exact match, respond with: "Please provide additional details."
            Make sure your answer is relevant and accurate to the question and does not repeat the question itself.
            In response give only answer but not question.
            """
            retriever = vectorstore.as_retriever(
                search_type="similarity",
                search_kwargs={
                    "k": 1,  
                    "filter": {
                        'tenantId': {'$eq': tenantId}
                    }
                }
            )
            docs = retriever.get_relevant_documents(keyword)
            exact_match = False
            exact_answer = None
            for doc in docs:
                if "Question:" in doc.page_content and "Answer:" in doc.page_content:
                    qa_parts = doc.page_content.split("\n")
                    if len(qa_parts) >= 2:
                        stored_question = qa_parts[0].replace("Question:", "").strip()
                        stored_answer = qa_parts[1].replace("Answer:", "").strip()
                        if stored_question.lower() == keyword.lower():
                            exact_match = True
                            exact_answer = stored_answer
                            break
            if exact_match:
                return PlainTextResponse(exact_answer)
            # else:
            #     prompt = ChatPromptTemplate.from_template(template)
            #     rag_chain = (
            #         {"context": retriever, "question": RunnablePassthrough()}
            #         | prompt
            #         | llm
            #         | StrOutputParser())
            #     answer = rag_chain.invoke(keyword)
            #     return PlainTextResponse(answer)

        else:
            raise HTTPException(status_code=400, detail="Invalid type parameter. Use 'upload', 'links', or 'qa'.")

    except Exception as e:
        error_msg = f"Error processing your request: {str(e)}. Please contact support for assistance."
        print(f"Error in unified_prompts_endpoint: {str(e)}") 
        return PlainTextResponse(error_msg, status_code=500)

############################ Summarizer #################################################
class SummarizationRequest(BaseModel):
    caseId: str
    assignedTo: str
    patientName: str
    onBehalfOf: str
    currentCaseStatus: str
    summarizationType: str
    notes: List[Any]

@app.post("/api/artificial-intelligence/summarize")
async def summarize(request: SummarizationRequest):
    """Combine JSON upload and summary generation into one endpoint."""
    try:
        json_data = {"caseId": request.caseId, "notes": request.notes, "patientName": request.patientName, "assignedTo": request.assignedTo, "currentCaseStatus": request.currentCaseStatus}
        noteLength = sum(len(note['note']) for note in request.notes)
        maxTokenToUse = 750
        
        if(noteLength <= 10):
            maxTokenToUse=200
        
        
        json_doc = Document(page_content=json.dumps(json_data), metadata={'fileName': 'query_json'})
        split_docs = [json_doc]
        print(f"Processing {len(split_docs)} documents.")
        
        llm = ChatOpenAI(
            model="gpt-4o-mini",
            temperature=0.0,
            max_tokens=maxTokenToUse,
            top_p=1.0,
        )

        full_summary_prompt_template = f"""
        "{{split_docs}}"

        Please summarize the following notes directly to me, the {request.onBehalfOf} in a structured format, focusing on key actions and communications. Include the following sections:
        
        1. Initial Request: Briefly explain the initial issue or request.
        2. Key Updates: Highlight major actions taken, communications, and progress in chronological order, including dates.
        3. Current Status: Summarize the present situation or outcome of the case.

        Do not expose any ids

        Complete the sentence at the end with a full stop.
        Give html format
        """

        unread_summary_prompt_template = f"""
        "{{split_docs}}"

        Please summarize the following case notes directly to me, the {request.onBehalfOf} in a paragraph format.
        
        Do not expose any ids
        Keep it simple and short.
        Complete the sentence at the end with a full stop.
        """


        prompt=""
        if(request.summarizationType == 'full'):
            prompt = full_summary_prompt_template
        else:
            prompt = unread_summary_prompt_template
        
        print(prompt)

        prompt = PromptTemplate.from_template(prompt)
        llm_chain = LLMChain(llm=llm, prompt=prompt)
        stuff_chain = StuffDocumentsChain(llm_chain=llm_chain, document_variable_name="split_docs")
        summary = stuff_chain.run(split_docs)

        sentences = summary
        print(sentences)

        # if len(sentences) > 1:
        #     summary = '.'.join(sentences[:-1]).strip()  
        # else:
        #     summary = sentences[0].strip()  

        # if not summary.endswith('.'):
        #     summary += '.'
        
        summary = sentences.replace("\n", " ")
        summary = summary.replace("```html", " ")
        return JSONResponse(content={
            "summary": summary
        })
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error during summarizing the documents: {str(e)}")

####################################### delete #########################
class DeleteRequest(BaseModel):
    prefix: str
@app.delete("/api/artificial-intelligence/delete", summary="Delete Documents", description="Delete all documents in the Pinecone index that match the given prefix.")
async def delete_documents(request: DeleteRequest):
    """Delete documents matching prefix
    Args:
        request: Delete request containing prefix
    Returns:
        JSON response indicating deletion status
    Raises:
        HTTPException: When no matching documents found
    """
    prefix = request.prefix
    ids_to_delete = [id for id in index.list(prefix=prefix)]
    print(ids_to_delete)
    if not ids_to_delete:
        raise HTTPException(status_code=404, detail="No documents found with the given prefix.")
    index.delete(ids=ids_to_delete)
    uploaded_documents_path = "do_not_delete_uploaded_documents.json"
    with open(uploaded_documents_path, "r") as f:
        upload_documents = json.load(f)
    if prefix in upload_documents:
        del upload_documents[prefix]
    with open(uploaded_documents_path, "w") as f:
        json.dump(upload_documents, f, indent=4)
    return {"message": f"Deleted {len(ids_to_delete)} documents with prefix '{prefix}'."}





if __name__ == "__main__":
    import uvicorn

    uvicorn.run(app, host="0.0.0.0", port=8000)
